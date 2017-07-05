# coding=utf-8
import docx
from docx.shared import Inches


doc = docx.Document(u'陕西联通集中优化周报_20170605-20170626_省公司.docx')

print(22)

for paragraph in  doc.paragraphs:
    print(paragraph.text)


#添加图片，设置图片大小
doc.add_picture(r"3G\\20170601_20170701_WCDMA_rnc异频硬切换成功率(%).png", width=Inches(2.25))













#保存文本
doc.save('demo.docx')